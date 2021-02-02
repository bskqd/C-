import openpyxl
from docx import Document
import re


BILL = r"(Акт\s+?№\s+?)([_]+)\b"
DATE = r"(Дата:?\s+?)(__.__.____)\b"
LIST = r"(Даний Акт засвідчує, що Виконавцем на майданчику:?\s+?)([_]+)\b"
LIST1 = r"(за адресом:?\s+?)([_]+)\b"
SUM = r"(Сума виконаних робіт складає:?\s+?)([_]+\s+)\b"
INFO = r"([_]+\s+)([_]+)\b"
SPACES = r"\s{84}"


def create_bill(act_no, data, template):
    wb = openpyxl.load_workbook(data)

    ws = wb["acts"]
    for row in ws.rows:
        if row[1].value == str(act_no):
            act_id, act_no, act_date, act_summa, ground_id = [c.value for c in row]
            break

    ws = wb["ground"]
    for row in ws.rows:
        if row[0].value == ground_id:
            ground_id, ground_name, ground_address, responsible, manager = [c.value for c in row]
            break

    works = []
    ws = wb["items"]
    for row in ws.rows:
        if row[1].value == act_id:
            work_id, act_id = [c.value for c in row]
            works.append(work_id)

    works_final = []
    ws = wb["work"]
    for row in ws.rows:
        work_id = row[0].value
        if work_id in works:
            work_id, name = [c.value for c in row]
            works_final.append(name)

    doc = Document(template)

    t = doc.tables[0]
    for i, work in enumerate(works_final, 1):

        values = (str(i), work)
        row = t.add_row()
        for cell, value in zip(row.cells, values):
            cell.text = value

    def change_act_no(match):
        return match.group(1) + str(act_no)

    def change_date(match):
        return match.group(1) + str(act_date)

    def change_ground(match):
        return match.group(1) + str(ground_name)

    def change_address(match):
        return match.group(1) + str(ground_address)

    def change_sum(match):
        return match.group(1) + str(act_summa)

    def change_responsible(match):
        return responsible + '                                             ' \
                             '                                             '\
               + manager

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = re.sub(BILL, change_act_no, run.text)
            run.text = re.sub(DATE, change_date, run.text)
            run.text = re.sub(LIST, change_ground, run.text)
            run.text = re.sub(LIST1, change_address, run.text)
            run.text = re.sub(SUM, change_sum, run.text)
            run.text = re.sub(INFO, change_responsible, run.text)

    doc.save("acts_" + ground_name + "__" + act_date + ".docx")


if __name__ == '__main__':
    create_bill(34, "data.xlsx", "Info.docx")
